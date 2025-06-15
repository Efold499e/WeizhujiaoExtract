import os 
try:
    import openpyxl
    import pandas as pd
    import numpy as np
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
except ImportError:
    os.system("pip install pandas numpy openpyxl python-docx")
    import pandas as pd
    import numpy as np
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn

class WeizhujiaoExtract:
    file_path = ""
    name = ""
    df = pd.DataFrame()
    answer_list = []
    question_list = []
    
    def __init__(self, file_path , name): # 如 马原05班课堂讨论(2025-05-29).xlsx , ***
        self.file_path = file_path
        self.name = name
        self.df = pd.read_excel(file_path)

    def get_question(self):# 得到问题
        question = ""
        question = self.df.columns[0]
        self.question_list.append(question[21:]) 
        print(question) 
        return question
    def get_answer(self):# 得到答案
        value = ""

        value = self.df.loc[self.df['Unnamed: 1'] == self.name, 'Unnamed: 6'].values # 直接找就好，微助教导出格式很唐
        if np.array_equal(value, np.array([], dtype=object)): # 如果没有找到答案，就返回空
            print("\n")
            self.answer_list.append("-1")
            return "-1"
        elif isinstance(value, int):
            return
        else:
            self.answer_list.append(value[0])
            print(value[0]+"\n") # 打印答案
            print(str(len(value[0]))+'\n')
            return value[0]
    def combine_qaa(self):# 同时得到问题和答案
        self.get_question()
        self.get_answer()

    def switch_sheet(self, function):# 这里function可以用上面三个中的任一
        sheet_names = pd.ExcelFile(self.file_path).sheet_names
        cnt = 0
        for sheet_name in sheet_names:
            cnt+=1
            if cnt == 1:
                continue
            df_sheet = pd.read_excel(self.file_path, sheet_name=sheet_name)
            self.df = df_sheet
            self.df.drop(index=0)
            function()
    
    def write_to_docx(self, file_name):# 将答案写入到docx文件中
        doc = Document()
        head = doc.add_paragraph(file_name[:-5])
        run = head.add_run()
        run.bold = True
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(14)
        i = 0
        cnt = 0
        for answer in self.answer_list:
            if (answer == "-1"):
                i+=1
                cnt+=1
                continue
            head = doc.add_paragraph(f"问题{i+1-cnt}：{self.question_list[i]}")
            run = head.add_run()
            run.bold = True
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(11)

            para = doc.add_paragraph(f'答案：{answer}')
            run = para.add_run()
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(11)
            i+=1
        doc.save(file_name)
        print(f"Answers saved to {file_name}")

if __name__ == '__main__':# 使用示例
    file_path = "马原07班课堂讨论(2025-05-29).xlsx"
    name = input("请输入姓名：")
    z = WeizhujiaoExtract(file_path, name)
    z.switch_sheet(z.combine_qaa)
    file_name = f"{name}马原讨论总结.docx"
    z.write_to_docx(file_name)


        
        